"""
DOCX Extractor Module - Handles Word documents (.docx)
"""

import os
import pandas as pd
from typing import Dict, List, Any, Union

from .base_extractor import BaseExtractor

# Import for Word documents
try:
    from docx import Document
    from docx.table import Table
except ImportError:
    print("Warning: python-docx not installed. DOCX support disabled.")
    Document = None
    Table = None


class DocxExtractor(BaseExtractor):
    """
    Extractor for Word documents (.docx)
    Extracts both text content and tables
    """
    
    def __init__(self):
        super().__init__()
        self.supported_extensions = ['.docx']
    
    def extract(self, file_path: str) -> Dict[str, Any]:
        """
        Extract text and tables from DOCX file
        """
        self.validate_file_exists(file_path)
        
        if not self.is_supported_file(file_path):
            raise ValueError(f"File extension not supported. Expected: {self.supported_extensions}")
        
        if Document is None:
            raise ImportError("python-docx not available. Install with: pip install python-docx")
        
        result = {
            'filename': os.path.basename(file_path),
            'file_type': 'docx',
            'text_content': '',
            'tables': [],
            'metadata': {}
        }
        
        try:
            doc = Document(file_path)
            
            # Extract paragraphs
            paragraphs = []
            for paragraph in doc.paragraphs:
                text = paragraph.text.strip()
                if text:
                    paragraphs.append(text)
            
            result['text_content'] = '\n\n'.join(paragraphs)
            
            # Extract tables
            tables = []
            for i, table in enumerate(doc.tables):
                table_data = self._extract_table_data(table, i + 1)
                tables.append(table_data)
            
            result['tables'] = tables
            result['metadata']['total_paragraphs'] = len(paragraphs)
            result['metadata']['total_tables'] = len(tables)
            
        except Exception as e:
            raise Exception(f"Error reading DOCX file: {str(e)}")
        
        return result
    
    def _extract_table_data(self, table: Any, table_index: int) -> Dict[str, Any]:
        """Extract data from a Word table"""
        if Table is None:
            return {
                'table_index': table_index,
                'shape': (0, 0),
                'has_header': False,
                'data': pd.DataFrame(),
                'raw_data': []
            }
        
        table_data = []
        
        for row in table.rows:
            row_data = []
            for cell in row.cells:
                # Clean cell text
                cell_text = cell.text.strip().replace('\n', ' ').replace('\r', '')
                row_data.append(cell_text)
            table_data.append(row_data)
        
        if not table_data:
            return {
                'table_index': table_index,
                'shape': (0, 0),
                'has_header': False,
                'data': pd.DataFrame(),
                'raw_data': []
            }
        
        # Normalize row lengths
        max_cols = max(len(row) for row in table_data)
        normalized_data = []
        for row in table_data:
            normalized_row = row + [''] * (max_cols - len(row))
            normalized_data.append(normalized_row)
        
        df = pd.DataFrame(normalized_data)
        has_header = self._detect_header_from_list(normalized_data)
        
        return {
            'table_index': table_index,
            'shape': (len(normalized_data), max_cols),
            'has_header': has_header,
            'data': df,
            'raw_data': normalized_data
        }
    
    def to_markdown(self, extracted_data: Dict[str, Any]) -> str:
        """
        Convert extracted data to markdown format - simplified, preserving natural structure
        """
        markdown_content = []
        
        # Add text content if it exists, preserving structure
        if extracted_data['text_content'].strip():
            markdown_content.append(extracted_data['text_content'])
            
            # Add spacing if there are tables after text
            if extracted_data['tables']:
                markdown_content.append("")
        
        # Add tables directly after text
        if extracted_data['tables']:
            for table in extracted_data['tables']:
                markdown_table = self.create_markdown_table_from_df(table['data'], table['has_header'])
                markdown_content.append(markdown_table)
                markdown_content.append("")
        
        return "\n".join(markdown_content).strip() 